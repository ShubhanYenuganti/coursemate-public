"""
Pass 1 extractor for DOCX files using python-docx.

Produces chunks of types: heading, paragraph, table, ocr.
page_number is always None (not natively available in DOCX format).

Embedded images are extracted from the word/media/ directory inside the
DOCX ZIP archive and passed through Tesseract OCR (via extract_image).
Images smaller than MIN_IMAGE_PIXELS are skipped to avoid noise from
icons and decorative elements.
"""
from docx import Document
from PIL import Image
from io import BytesIO
from typing import List, Dict, Any
import zipfile
from .image import extract_image

# Skip images whose total pixel area is below this threshold
MIN_IMAGE_PIXELS = 100 * 100  # 10 000 px

# Only OCR raster formats; skip vector/audio/other media in the archive
OCR_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif"}


def _extract_text(file_bytes: bytes) -> List[Dict[str, Any]]:
    doc = Document(BytesIO(file_bytes))
    chunks: List[Dict[str, Any]] = []

    # ── paragraphs ────────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "") if para.style else ""
        chunk_type = "heading" if style_name.startswith("Heading") else "paragraph"

        chunks.append({
            "text": text,
            "chunk_type": chunk_type,
            "page_number": None,
            "token_count": len(text.split()),
        })

    # ── tables ────────────────────────────────────────────────────────────────
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells: List[str] = []
            for header, cell in zip(headers, row.cells):
                v = cell.text.strip()
                if v:
                    cells.append(f"{header}: {v}" if header else v)
            row_text = " | ".join(cells)
            if row_text.strip():
                chunks.append({
                    "text": row_text,
                    "chunk_type": "table",
                    "page_number": None,
                    "token_count": len(row_text.split()),
                })

    return chunks


def _extract_embedded_images(file_bytes: bytes) -> List[Dict[str, Any]]:
    """
    Pull every raster image from the word/media/ directory inside the DOCX
    ZIP archive and run OCR on each one that passes the size threshold.
    """
    chunks: List[Dict[str, Any]] = []
    try:
        with zipfile.ZipFile(BytesIO(file_bytes)) as z:
            media_entries = [
                name for name in z.namelist()
                if name.startswith("word/media/")
                and any(name.lower().endswith(ext) for ext in OCR_EXTENSIONS)
            ]
            for entry in media_entries:
                try:
                    img_bytes = z.read(entry)
                    with Image.open(BytesIO(img_bytes)) as im:
                        w, h = im.size
                    if w * h < MIN_IMAGE_PIXELS:
                        continue
                    chunks.extend(extract_image(img_bytes))
                except Exception:
                    pass  # one bad image must not abort the rest
    except zipfile.BadZipFile:
        pass  # not a valid DOCX ZIP — shouldn't happen, but be safe
    return chunks


def extract_docx(file_bytes: bytes) -> List[Dict[str, Any]]:
    return _extract_text(file_bytes) + _extract_embedded_images(file_bytes)
